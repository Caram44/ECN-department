#use this program to create output for scholarships
#the  file contains all courses taken by ECN majors
#use virtualenv awards.venv
import sys
sys.path.append('/Users/cmcdanie/ECNdept/dept_python/library/')
import pandas as pd
from docx import Document
from award_functions import subject_GPA, list_courses_noprof_docx
from award_functions import drop_dup, make_header_docx
from award_functions import list_courses_prof_docx, student_report
document = Document()

tbl = pd.read_csv('/Users/cmcdanie/ECN_dept_data/input/ECN_major_classes_2171.csv',
                  parse_dates=True)

#Rondthaler conditions
#Economics major; Junior in standing; highest cumulative GPA;
#highest GPA in economics;
#completed 12 hours of upper-division economics courses @ ASU Tempe Campus;
#most course work at ASU with a min. of 30 hours @ Tempe Campus;
#strength of course work.

#condition on level, GPA and major -
#Rondthaler required to be WPC
tbl = tbl[tbl['Acad Plan'] == 'BAECNBS']

tbl = tbl[tbl['Cum Gpa'] > 3.75]
tbl = tbl.sort(['Cum Gpa'], ascending=False)
tbl['Catalog Nbr'] = tbl['Catalog Nbr'].convert_objects(convert_numeric=True)

#Rondthaler scholars
micro = tbl[tbl['Catalog Nbr'].isin([214, 312])]['Emplid']
macro = tbl[tbl['Catalog Nbr'].isin([213, 313])]['Emplid']
both = macro[macro.isin(micro) == True]


#students to create schedules
IDs = both.unique()
email = []
count = 0
for ID in IDs:
    print(ID)
    df = student_report(ID, tbl)
    df = drop_dup(df)
    #count upper-division ECN courses
    ECNs = df[df['Subject'] == 'ECN']
    level = ECNs['Catalog Nbr'].convert_objects(convert_numeric=True)
    ECNud = len(level[level > 300])
    if ECNud > 3:
        ECN_GPA = subject_GPA('ECN', df)
        #Create a word table - yes it sucks
        table = document.add_table(rows=1, cols=5, style='MediumList2')
        hdr_cells = table.rows[0].cells
        make_header_docx(df, document, hdr_cells)
        list_courses_prof_docx(ECNs, document, table)
        MATs = df[df['Subject'].isin(['MAT', 'STP'])]
        list_courses_noprof_docx(MATs, document, table)
        paragraph = document.add_paragraph(' ')
        if count % 2 == 1:
                document.add_page_break()
        count = count + 1
document.save('/Users/cmcdanie/ECN_dept_data/output/awards/Rondthaler_2171.docx')
